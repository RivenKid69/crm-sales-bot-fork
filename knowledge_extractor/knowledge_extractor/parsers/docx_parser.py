"""DOCX document parser using python-docx."""

import logging
from pathlib import Path
from typing import List

from .base import AbstractParser, ContentType, ParsedDocument

logger = logging.getLogger(__name__)


class DOCXParser(AbstractParser):
    """Parse DOCX documents using python-docx."""

    SUPPORTED_EXTENSIONS = {".docx"}

    def supports(self, filepath: Path) -> bool:
        return filepath.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def parse(self, filepath: Path) -> ParsedDocument:
        """Parse DOCX and extract text with structure."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        filepath = Path(filepath)
        logger.info(f"Parsing DOCX: {filepath}")

        text_parts = []
        sections = []
        tables = []
        current_section = []

        try:
            doc = Document(filepath)

            # Extract metadata
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
            }

            # Process paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                text_parts.append(text)

                # Check if heading
                if para.style and para.style.name.startswith("Heading"):
                    if current_section:
                        sections.append("\n".join(current_section))
                        current_section = []

                current_section.append(text)

            # Add last section
            if current_section:
                sections.append("\n".join(current_section))

            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)

        except Exception as e:
            logger.error(f"Error parsing DOCX {filepath}: {e}")
            raise

        raw_text = self.clean_text("\n".join(text_parts))

        # Determine content type
        content_type = ContentType.PROSE
        if tables and len(tables) > len(sections):
            content_type = ContentType.TABLE
        elif self._has_qa_structure(raw_text):
            content_type = ContentType.QA_PAIRS

        return ParsedDocument(
            source_file=str(filepath),
            content_type=content_type,
            raw_text=raw_text,
            metadata=metadata,
            language=self.detect_language(raw_text),
            sections=sections if sections else [raw_text],
            tables=tables,
        )

    def _has_qa_structure(self, text: str) -> bool:
        """Check if text has Q&A structure."""
        import re

        # Look for Q: A: or Вопрос: Ответ: patterns
        qa_patterns = [
            r"(?:Q|В|Вопрос)[:\s]+.*?(?:A|О|Ответ)[:\s]+",
            r"\?\s*\n+[А-Яа-яA-Za-z]",  # Question mark followed by answer
        ]

        for pattern in qa_patterns:
            if len(re.findall(pattern, text, re.IGNORECASE)) > 3:
                return True
        return False
